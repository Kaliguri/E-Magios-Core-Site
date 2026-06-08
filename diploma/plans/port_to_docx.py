# -*- coding: utf-8 -*-
"""Порт манускрипта (Markdown) в новый .docx на основе оригинала-шаблона.
Сохраняет ГОСТ-стили оригинала, заменяет тело с «Аннотации» новым текстом."""
import re, copy, struct, os
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR = r"C:\Gamedev\E-Magios-Core-Site\diploma\figures"
CONTENT_W_MM = 165.0  # A4 210 - левое 30 - правое 15

def png_size(path):
    with open(path, 'rb') as f:
        head = f.read(24)
    if head[:8] != b'\x89PNG\r\n\x1a\n':
        return None
    w, h = struct.unpack('>II', head[16:24])
    return w, h

SRC = r"C:\Gamedev\E-Magios-Core-Site\diploma\ВКР_221-3711_ГайдарьМД.05.05.docx"
MD  = r"C:\Gamedev\E-Magios-Core-Site\diploma\plans\zapiska_manuscript.md"
OUT = r"C:\Gamedev\E-Magios-Core-Site\diploma\v1\Пояснительная записка 221-3711 Гайдарь М.Д.docx"

doc = Document(SRC)

# --- ГОСТ-поля: правое 15 мм (в оригинале 10), левое 30, верх/низ 20 ------
for sec in doc.sections:
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(15)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)

# --- 0. Правки титула / задания / календарного плана ---------------------
def set_para_text(p, text):
    """Переписать текст абзаца, сохранив форматирование первого run."""
    runs = p.runs
    if not runs:
        p.add_run(text); return
    runs[0].text = text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)

def clear_numbering(p):
    pPr = p._p.get_or_add_pPr()
    for numPr in pPr.findall(qn('w:numPr')):
        pPr.remove(numPr)

def set_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    set_para_text(p, text)
    clear_numbering(p)
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

def fix_front_matter():
    centered_lines = (
        'МИНИСТЕРСТВО НАУКИ',
        'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ',
        'МОСКОВСКИЙ ПОЛИТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ',
        'Факультет Информационных технологий',
        'кафедра Информатики и информационных технологий',
        'направление подготовки 09.03.02',
    )
    for p in doc.paragraphs:
        if p.text.strip().lower() == 'аннотация':
            break
        t = p.text
        stripped = t.strip()
        if any(stripped.startswith(line) for line in centered_lines):
            if stripped.startswith('направление подготовки 09.03.02'):
                set_para_text(p, 'направление подготовки 09.03.02 «Информационные системы и технологии»')
            elif stripped.startswith('кафедра Информатики и информационных технологий'):
                set_para_text(p, 'кафедра Информатики и информационных технологий')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
        # шифр направления и опечатка в отчестве
        for r in p.runs:
            if '09.002' in r.text:
                r.text = r.text.replace('09.002', '09.03.02')
            if 'Деинсович' in r.text:
                r.text = r.text.replace('Деинсович', 'Денисович')
        # заполнение полей задания
        if t.strip().startswith('Утверждена приказом'):
            set_para_text(p, 'Утверждена приказом по Университету '
                             '№ 7174-УД от «30» декабря 2025 г.')
        elif t.strip().startswith('Срок представления работы к защите'):
            set_para_text(p, 'Срок представления работы к защите '
                             '« 09 » июня 2026 г.')
        elif t.strip().startswith('Дата выдачи задания'):
            set_para_text(p, 'Дата выдачи задания: «25» сентября 2025 г.')

    # календарный план: нумерация этапов 1–8 и согласованные даты
    stage_dates = {
        'Исследование предметной области': '26.09.2025-19.10.2025',
        'Разработка технического задания': '20.10.2025-02.11.2025',
        'Эскизное проектирование': '03.11.2025-23.11.2025',
        'Техническое проектирование': '24.11.2025-21.12.2025',
        'Рабочее проектирование': '22.12.2025-01.02.2026',
        'Разработка': '02.02.2026-29.03.2026',
        'Тестирование': '30.03.2026-26.04.2026',
        'Приемосдаточные испытания': '27.04.2026-24.05.2026',
    }
    for tbl in doc.tables:
        n = 0
        for row in tbl.rows:
            cells = row.cells
            if len(cells) < 3:
                continue
            name = cells[1].text.strip()
            if name in stage_dates:
                n += 1
                set_cell_text(cells[0], f'{n}.')
                set_cell_text(cells[2], stage_dates[name])

fix_front_matter()

# --- 1. Удалить тело начиная с «Аннотация» -------------------------------
body = doc.element.body
split_el = None
for p in doc.paragraphs:
    if p.text.strip().lower() == 'аннотация':
        split_el = p._element
        break
if split_el is None:
    raise SystemExit("Не найден разделитель 'Аннотация'")
# собрать все элементы после (и включая) split_el на уровне body и удалить,
# кроме финального sectPr
to_remove = []
hit = False
for child in list(body):
    if child is split_el:
        hit = True
    if hit and child.tag != qn('w:sectPr'):
        to_remove.append(child)
for el in to_remove:
    body.remove(el)

# финальный sectPr (последний элемент body) — точка вставки нового контента
sectPr = body.find(qn('w:sectPr'))

def add_para():
    p = OxmlElement('w:p')
    if sectPr is not None:
        sectPr.addprevious(p)
    else:
        body.append(p)
    from docx.text.paragraph import Paragraph
    return Paragraph(p, doc)

# --- helpers -------------------------------------------------------------
def set_border(par):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6')
        e.set(qn('w:space'),'6'); e.set(qn('w:color'),'808080')
        pbdr.append(e)
    pPr.append(pbdr)

def set_shading(par, fill='F0F0F0'):
    pPr = par._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),fill)
    pPr.append(sh)

def no_indent(par):
    par.paragraph_format.first_line_indent = Pt(0)
    par.paragraph_format.left_indent = Pt(0)

REF = re.compile(r'\[ССЫЛКА:(\d+)\]')
INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+?`)')

def add_runs(par, text, bold=False, italic=False, mono=False, size=None):
    text = REF.sub(r'[\1]', text)
    for seg in INLINE.split(text):
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            r = par.add_run(seg[2:-2]); r.bold = True
        elif seg.startswith('`') and seg.endswith('`'):
            r = par.add_run(seg[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(11)
        else:
            r = par.add_run(seg); r.bold = bold; r.italic = italic
            if mono: r.font.name = 'Consolas'
        if size: r.font.size = size

def heading(text, level, struct=False, page_break=True, in_toc=True):
    """level: 1 глава/структ, 2 раздел, 3 подраздел. struct=по центру прописными."""
    p = add_para()
    if in_toc:
        p.style = doc.styles[f'Heading {level}']
    else:
        p.style = doc.styles['Normal']
    p.paragraph_format.page_break_before = page_break
    no_indent(p)
    if struct:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = True
        if not in_toc:
            r.font.size = Pt(14)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text); r.bold = True
    return p

def toc_field():
    h = heading('СОДЕРЖАНИЕ', 1, struct=True, page_break=True, in_toc=False)
    p = add_para()
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u \b TOCSCOPE')
    ph = OxmlElement('w:p'); rr = OxmlElement('w:r'); t = OxmlElement('w:t')
    t.text = 'Оглавление обновится при открытии в Word (F9).'
    rr.append(t); ph.append(rr); fld.append(ph)
    p._p.append(fld)

def figure(num, desc):
    img = os.path.join(FIG_DIR, 'fig_' + num.replace('.', '_') + '.png')
    pic = add_para(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_indent(pic)
    if os.path.exists(img):
        size = png_size(img)
        if size:
            w_px, h_px = size
            aspect = w_px / h_px
            # ГОСТ: высота рисунка не больше ~половины страницы (поле ~257 мм),
            # и не увеличивать растр сверх натурального размера (≈96 dpi).
            MAX_H_MM = 118.0
            natural_w_mm = w_px / 96.0 * 25.4
            w_mm = min(CONTENT_W_MM, MAX_H_MM * aspect, natural_w_mm)
            pic.add_run().add_picture(img, width=Mm(w_mm))
        else:
            pic.add_run().add_picture(img, width=Mm(CONTENT_W_MM))
    else:
        set_border(pic); set_shading(pic, 'EAEAEA')
        r = pic.add_run(f'[ Рисунок {num} — скриншот будет вставлен ]')
        r.bold = True; r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    cap = add_para(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_indent(cap)
    cap.add_run(f'Рисунок {num} — {desc}')

def table_caption(num, title):
    cap = add_para(); cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT; no_indent(cap)
    cap.add_run(f'Таблица {num} — {title}')

def set_table_borders(t):
    tblPr = t._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')
        borders.append(e)
    tblPr.append(borders)

def render_table(rows):
    """rows: список списков ячеек (первая строка — заголовок)."""
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    set_table_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # ширины колонок: первая шире
    total = CONTENT_W_MM
    first = total * 0.34 if ncol >= 3 else total / ncol
    rest = (total - first) / (ncol - 1) if ncol > 1 else total
    widths = [first] + [rest] * (ncol - 1)
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = t.cell(ri, ci)
            cell.width = Mm(widths[ci])
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            txt = row[ci] if ci < len(row) else ''
            r = p.add_run(txt); r.font.size = Pt(12)
            if ri == 0:
                r.bold = True
    # переместить таблицу перед финальным sectPr
    if sectPr is not None:
        sectPr.addprevious(t._tbl)
    # отбивка снизу
    sp = add_para(); sp.paragraph_format.first_line_indent = Pt(0)
    sp.add_run('')

def table_placeholder(num, desc):
    cap = add_para(); cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT; no_indent(cap)
    cap.add_run(f'Таблица {num} — {desc.split("—")[0].strip()[:60]}')
    box = add_para(); box.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_indent(box)
    set_border(box); set_shading(box, 'F4F4F4')
    box.add_run(f'[ Таблица {num}: {desc} ]').italic = True

def listing_caption(text):
    p = add_para(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; no_indent(p)
    p.add_run(text)

def code_block(lines):
    p = add_para(); no_indent(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_border(p); set_shading(p, 'F6F6F6')
    for i, ln in enumerate(lines):
        if i: p.add_run().add_break()
        r = p.add_run(ln if ln else ' '); r.font.name = 'Consolas'; r.font.size = Pt(10)

def appendix_code_block(lines):
    """Листинг приложения: построчно, без рамки — чтобы переносился по страницам."""
    for ln in lines:
        p = add_para(); no_indent(p); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(ln if ln else ' '); r.font.name = 'Consolas'; r.font.size = Pt(10)

def slides():
    sl_dir = os.path.join(FIG_DIR, 'slides')
    files = sorted(f for f in os.listdir(sl_dir) if f.lower().endswith('.png'))
    for i, f in enumerate(files, 1):
        pic = add_para(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_indent(pic)
        pic.paragraph_format.space_before = Pt(6)
        pic.add_run().add_picture(os.path.join(sl_dir, f), width=Mm(150))
        cap = add_para(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_indent(cap)
        r = cap.add_run('Слайд %d' % i); r.font.size = Pt(12)

def normal(text):
    p = add_para(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)

def list_item(text, dash=True):
    p = add_para(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(18)
    add_runs(p, text)

# --- 2. Парсинг манускрипта ---------------------------------------------
raw = open(MD, encoding='utf-8').read()
# начать с первой структурной шапки
start = raw.index('# Аннотация')
lines = raw[start:].split('\n')

STRUCT_NO_TOC = {'АННОТАЦИЯ','ABSTRACT','РЕФЕРАТ'}
STRUCT_TOC = {'ВВЕДЕНИЕ','ЗАКЛЮЧЕНИЕ','БИБЛИОГРАФИЧЕСКИЙ СПИСОК'}

i = 0
fig_re = re.compile(r'^\[РИСУНОК\s+([^\]:]+):\s*(.+?)\]$')
tab_re = re.compile(r'^\[ТАБЛИЦА\s+([^\]:]+):\s*(.+?)\]$')
lst_re = re.compile(r'^\[ЛИСТИНГ\s+(.+?)\]$')
toc_inserted = False
intro_para = None
in_appendix = False

while i < len(lines):
    line = lines[i].rstrip()
    s = line.strip()
    if not s or s == '---':
        i += 1; continue

    if s.startswith('# '):
        title = s[2:].strip()
        up = title.upper()
        if up.startswith('ПРИЛОЖЕНИЕ'):
            in_appendix = True
            p = add_para(); p.style = doc.styles['Heading 1']
            p.paragraph_format.page_break_before = True
            no_indent(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(title).bold = True
            i += 1; continue
        if up.startswith('ГЛАВА'):
            heading(title, 1, struct=False)
        elif up in STRUCT_NO_TOC:
            heading(title, 1, struct=True, in_toc=False)
        elif up in STRUCT_TOC:
            if up == 'ВВЕДЕНИЕ' and not toc_inserted:
                toc_field(); toc_inserted = True
            p = heading(title, 1, struct=True, in_toc=True)
            if up == 'ВВЕДЕНИЕ':
                bm = OxmlElement('w:bookmarkStart')
                bm.set(qn('w:id'), '900'); bm.set(qn('w:name'), 'TOCSCOPE')
                p._p.insert(0, bm)
                intro_para = p
                p.paragraph_format.page_break_before = False  # новую страницу даст разрыв раздела
        else:
            heading(title, 1, struct=True, in_toc=False)
        i += 1; continue

    if s.startswith('## '):
        heading(s[3:].strip(), 2, struct=False, page_break=False, in_toc=not in_appendix)
        i += 1; continue
    if s.startswith('### '):
        heading(s[4:].strip(), 3, struct=False, page_break=False, in_toc=not in_appendix)
        i += 1; continue

    if s == '[СЛАЙДЫ]':
        slides(); i += 1; continue

    m = fig_re.match(s)
    if m:
        figure(m.group(1), m.group(2)); i += 1; continue
    m = tab_re.match(s)
    if m:
        table_caption(m.group(1), m.group(2)); i += 1; continue
    m = lst_re.match(s)
    if m:
        listing_caption('Листинг ' + m.group(1).replace(':', ' —', 1)); i += 1; continue

    if s.startswith('```'):
        block = []; i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            block.append(lines[i]); i += 1
        (appendix_code_block if in_appendix else code_block)(block); i += 1; continue

    if s.startswith('|'):
        block = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            block.append(lines[i].strip()); i += 1
        rows = []
        for bl in block:
            cells = [c.strip() for c in bl.strip('|').split('|')]
            if all(c and set(c) <= set('-: ') for c in cells):
                continue  # строка-разделитель |---|
            rows.append(cells)
        if rows:
            render_table(rows)
        continue

    if s.startswith('- '):
        list_item('– ' + s[2:].strip()); i += 1; continue
    m = re.match(r'^(\d+)\.\s+(.+)$', s)
    if m:
        list_item(f'{m.group(1)}) {m.group(2)}'); i += 1; continue

    normal(s); i += 1

# --- закрыть закладку TOCSCOPE перед финальным sectPr -------------------
bm_end = OxmlElement('w:bookmarkEnd'); bm_end.set(qn('w:id'), '900')
if sectPr is not None:
    sectPr.addprevious(bm_end)
else:
    body.append(bm_end)

# --- разрыв раздела перед «Введением» + номера страниц с «Введения» ------
if intro_para is not None and sectPr is not None:
    brk = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sect_copy = copy.deepcopy(sectPr)
    for ref in sect_copy.findall(qn('w:footerReference')) + sect_copy.findall(qn('w:headerReference')):
        sect_copy.remove(ref)
    pPr.append(sect_copy)
    brk.append(pPr)
    intro_para._p.addprevious(brk)
    # чистый лист: убрать футеры и перезапуски нумерации из ВСЕХ разделов,
    # чтобы нумерация была сквозной физической (без рестарта на Аннотации)
    for sp in body.iter(qn('w:sectPr')):
        for fr in sp.findall(qn('w:footerReference')):
            sp.remove(fr)
        for pnt in sp.findall(qn('w:pgNumType')):
            sp.remove(pnt)
    # Введение и далее — отдельный раздел С НОВОЙ СТРАНИЦЫ (в шаблоне был
    # continuous, из-за чего Введение «прилипало» к Определениям)
    sect_type = sectPr.find(qn('w:type'))
    if sect_type is None:
        sect_type = OxmlElement('w:type'); sectPr.insert(0, sect_type)
    sect_type.set(qn('w:val'), 'nextPage')
    # основной раздел (Введение→конец): номер страницы по центру снизу;
    # остальные разделы остаются без футера → без номеров
    main = doc.sections[-1]
    main.footer.is_linked_to_previous = False
    fp = main.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.first_line_indent = Pt(0)
    run = fp.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)

# --- 3. updateFields, чтобы Word обновил оглавление при открытии ---------
settings = doc.settings.element
if settings.find(qn('w:updateFields')) is None:
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'),'true')
    settings.insert(0, uf)

doc.save(OUT)
print("OK saved:", OUT)
print("Параграфов в новом документе:", len(doc.paragraphs))
